from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = "Курсовая_работа_GasServiceApp.docx"


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name, size in [("Heading 1", 14), ("Heading 2", 14), ("Heading 3", 14)]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER if style_name == "Heading 1" else WD_ALIGN_PARAGRAPH.LEFT

    footer = section.footer.paragraphs[0]
    add_page_number(footer)
    return doc


def p(doc, text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=True, bold=False):
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.first_line_indent = Cm(1.25) if first_indent else Cm(0)
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.bold = bold
    return para


def heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def caption(doc, text):
    para = p(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
    for run in para.runs:
        run.font.size = Pt(12)
    return para


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        set_cell_text(hdr.cells[i], h, bold=True)
        set_cell_shading(hdr.cells[i], "D9EAF7")
        if widths:
            hdr.cells[i].width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
            if widths:
                cells[i].width = Cm(widths[i])
    doc.add_paragraph()
    return table


def add_code_block(doc, title, code):
    caption(doc, title)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F2F2F2")
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.first_line_indent = Cm(0)
    para.paragraph_format.line_spacing = 1.0
    for idx, line in enumerate(code.strip().splitlines()):
        if idx:
            para.add_run().add_break()
        run = para.add_run(line)
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.size = Pt(9)
    doc.add_paragraph()


def add_title_page(doc):
    for text in [
        "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ",
        "Наименование образовательной организации",
        "Кафедра информационных систем и технологий",
    ]:
        p(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
    p(doc)
    p(doc)
    title = p(doc, "КУРСОВАЯ РАБОТА", align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False, bold=True)
    title.runs[0].font.size = Pt(16)
    p(doc, "по дисциплине: «Разработка программных модулей»", align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
    topic = p(doc, "на тему: «Разработка информационной системы обслуживания газового оборудования»", align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False, bold=True)
    topic.runs[0].font.size = Pt(16)
    for _ in range(5):
        p(doc)
    p(doc, "Выполнил: студент группы ____________ ____________________", align=WD_ALIGN_PARAGRAPH.RIGHT, first_indent=False)
    p(doc, "Руководитель: ____________________ ____________________", align=WD_ALIGN_PARAGRAPH.RIGHT, first_indent=False)
    for _ in range(6):
        p(doc)
    p(doc, "Саратов 2026", align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
    doc.add_page_break()


def add_contents(doc):
    heading(doc, "СОДЕРЖАНИЕ", 1)
    entries = [
        ("Введение", "3"),
        ("1 Аналитическая часть", "5"),
        ("1.1 Характеристика предметной области", "5"),
        ("1.2 Анализ требований к программному продукту", "7"),
        ("2 Проектная часть", "10"),
        ("2.1 Архитектура приложения", "10"),
        ("2.2 Проектирование базы данных", "12"),
        ("2.3 Проектирование пользовательского интерфейса", "15"),
        ("3 Реализация программного продукта", "17"),
        ("3.1 Средства разработки", "17"),
        ("3.2 Описание основных модулей", "19"),
        ("3.3 Работа с базой данных", "22"),
        ("4 Тестирование и эксплуатация", "24"),
        ("Заключение", "27"),
        ("Список использованных источников", "28"),
        ("Приложение А. Фрагменты программного кода", "30"),
        ("Приложение Б. SQL-скрипт структуры базы данных", "32"),
    ]
    for name, page in entries:
        line = f"{name}{'.' * max(2, 78 - len(name))}{page}"
        p(doc, line, align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=False)
    doc.add_page_break()


def build_body(doc):
    heading(doc, "ВВЕДЕНИЕ", 1)
    intro = [
        "Современные организации, выполняющие обслуживание газового оборудования, работают с большим объемом эксплуатационных данных: сведениями о клиентах, перечнем установленного оборудования, заявками на диагностику и техническое обслуживание, результатами выполненных работ. Ведение такой информации в разрозненных таблицах или бумажных журналах снижает скорость обработки заявок и повышает риск ошибок при планировании работ.",
        "Актуальность темы курсовой работы определяется необходимостью разработки простой и расширяемой информационной системы, которая позволит хранить сведения о газовом оборудовании и клиентах в централизованной базе данных, а также предоставит пользователю удобный настольный интерфейс для просмотра данных.",
        "Цель работы - разработать программное приложение «Система обслуживания газового оборудования», обеспечивающее учет оборудования и клиентов с использованием настольного интерфейса и реляционной базы данных PostgreSQL.",
        "Для достижения цели были поставлены задачи: изучить предметную область; определить функциональные и нефункциональные требования; спроектировать структуру базы данных; разработать модель данных приложения; реализовать подключение к PostgreSQL; создать пользовательский интерфейс; выполнить проверку сборки и работоспособности основных модулей.",
        "Объектом исследования является процесс учета оборудования и клиентов газовой сервисной организации. Предметом исследования являются программные средства автоматизации учета и просмотра эксплуатационных данных.",
        "Практическая значимость работы заключается в том, что разработанное приложение может использоваться как основа для дальнейшего развития системы: добавления регистрации заявок, фильтрации, поиска, формирования отчетов и разграничения прав доступа.",
    ]
    for text in intro:
        p(doc, text)

    heading(doc, "1 АНАЛИТИЧЕСКАЯ ЧАСТЬ", 1)
    heading(doc, "1.1 Характеристика предметной области", 2)
    for text in [
        "Газовое оборудование относится к объектам, требующим регулярного контроля технического состояния. Для сервисной организации важны точность сведений об оборудовании, история заявок, сроки обслуживания, данные ответственных клиентов и возможность оперативно определить текущий статус объекта.",
        "В рамках курсовой работы рассматривается упрощенная модель деятельности сервисной организации. Система хранит данные о клиентах, оборудовании и заявках на обслуживание. Клиент может иметь несколько обращений, а заявка связывается с конкретной единицей оборудования. Такая структура позволяет в дальнейшем строить отчеты по активности клиентов, состоянию оборудования и загрузке специалистов.",
        "Основными пользователями системы являются диспетчер или инженер сервисной службы. Пользователь должен видеть перечень оборудования, контролировать его состояние и иметь доступ к контактным данным клиента. Для начальной версии приложения достаточно режима просмотра данных, так как первоочередной задачей является надежная интеграция интерфейса с базой данных.",
    ]:
        p(doc, text)

    caption(doc, "Рисунок 1 - Обобщенная схема предметной области")
    add_table(
        doc,
        ["Сущность", "Назначение", "Связи"],
        [
            ["Клиент", "Организация или физическое лицо, владеющее оборудованием или обращающееся за обслуживанием", "Связан с заявками"],
            ["Оборудование", "Газовый котел, регулятор давления, водонагреватель или другой объект учета", "Связано с заявками"],
            ["Заявка", "Факт обращения или планового обслуживания", "Связана с клиентом и оборудованием"],
        ],
        [3.2, 9.5, 4.5],
    )

    heading(doc, "1.2 Анализ требований к программному продукту", 2)
    for text in [
        "Функциональные требования описывают действия, которые должна выполнять система. В разработанном приложении реализованы загрузка списка оборудования из базы данных, загрузка списка клиентов, отображение данных в табличном виде и вывод сообщения об ошибке при невозможности подключения к PostgreSQL.",
        "Нефункциональные требования определяют качество работы программы. Приложение должно запускаться на рабочей станции пользователя, иметь понятный интерфейс, не блокировать окно во время загрузки данных и использовать типизированный доступ к базе данных. Для этого в проекте применяется асинхронная загрузка данных и Entity Framework Core.",
        "Система разрабатывалась как учебный программный продукт, поэтому в ней сознательно сохранена простая архитектура без избыточной инфраструктуры. Это облегчает сопровождение кода и демонстрирует базовые принципы построения настольного приложения с подключением к реляционной СУБД.",
    ]:
        p(doc, text)
    add_table(
        doc,
        ["Требование", "Описание", "Реализация в проекте"],
        [
            ["Просмотр оборудования", "Пользователь видит список установленного оборудования", "Вкладка «Оборудование», DataGrid, EquipmentList"],
            ["Просмотр клиентов", "Пользователь видит список клиентов и контактные сведения", "Вкладка «Клиенты», DataGrid, ClientList"],
            ["Работа с PostgreSQL", "Данные хранятся в реляционной базе данных", "GasDbContext, Npgsql.EntityFrameworkCore.PostgreSQL"],
            ["Обработка ошибки подключения", "Пользователь получает сообщение при недоступной БД", "StatusMessage в MainViewModel"],
        ],
        [3.5, 6.7, 7.0],
    )

    heading(doc, "2 ПРОЕКТНАЯ ЧАСТЬ", 1)
    heading(doc, "2.1 Архитектура приложения", 2)
    for text in [
        "Приложение разработано на платформе .NET 8 с использованием фреймворка Avalonia UI. Avalonia позволяет создавать кроссплатформенные настольные интерфейсы с XAML-разметкой, а также применять подход MVVM, при котором представление отделяется от логики получения и подготовки данных.",
        "В проекте выделены следующие логические уровни: уровень представления, уровень модели представления, сервисный уровень, уровень доступа к данным и уровень доменных моделей. Такое разделение делает код более понятным и уменьшает связанность между интерфейсом и базой данных.",
        "Представление MainWindow.axaml описывает окно и вкладки интерфейса. MainViewModel хранит списки данных для отображения и статусное сообщение. GasService инкапсулирует операции получения и добавления сущностей. GasDbContext настраивает соединение с PostgreSQL и правила сопоставления классов с таблицами.",
    ]:
        p(doc, text)

    caption(doc, "Рисунок 2 - Логическая архитектура приложения")
    add_table(
        doc,
        ["Уровень", "Файлы проекта", "Ответственность"],
        [
            ["Представление", "Views/MainWindow.axaml", "Отображение вкладок, таблиц и статусного сообщения"],
            ["Модель представления", "ViewModels/MainViewModel.cs", "Загрузка данных, свойства привязки, обработка ошибок"],
            ["Сервис", "Services/GasService.cs", "Методы работы с клиентами и оборудованием"],
            ["Доступ к данным", "Services/GasDbContext.cs", "Настройка Entity Framework Core и PostgreSQL"],
            ["Модели", "Models/*.cs", "Описание сущностей предметной области"],
        ],
        [3.5, 5.8, 7.9],
    )

    heading(doc, "2.2 Проектирование базы данных", 2)
    for text in [
        "База данных gas_service_app содержит три основные таблицы: Equipment, Clients и ServiceRequests. Таблица Equipment предназначена для хранения сведений об оборудовании, включая серийный номер, тип, производителя, модель, место установки и статус. Таблица Clients хранит наименование клиента, контактное лицо, телефон и адрес.",
        "Таблица ServiceRequests фиксирует заявки на обслуживание. Она содержит ссылки на оборудование и клиента, дату поступления заявки, дату завершения, статус и описание работ. Внешние ключи настроены с ограничением ON DELETE RESTRICT, что предотвращает случайное удаление клиента или оборудования, если на них ссылаются заявки.",
        "Для обеспечения уникальности серийных номеров оборудования создан уникальный индекс IX_Equipment_SerialNumber. Для ускорения выборок по заявкам созданы индексы по внешним ключам EquipmentId и ClientId.",
    ]:
        p(doc, text)
    add_table(
        doc,
        ["Таблица", "Ключевые поля", "Назначение"],
        [
            ["Equipment", "Id, SerialNumber, Type, InstallationDate, Status", "Учет газового оборудования"],
            ["Clients", "Id, Name, ContactPerson, Phone, Address", "Учет клиентов сервисной организации"],
            ["ServiceRequests", "Id, EquipmentId, ClientId, RequestDate, Status", "Учет заявок и выполненных работ"],
        ],
        [4.0, 7.0, 7.2],
    )
    caption(doc, "Рисунок 3 - ER-модель базы данных")
    add_table(
        doc,
        ["Связь", "Кардинальность", "Смысл"],
        [
            ["Clients -> ServiceRequests", "1:N", "Один клиент может иметь несколько заявок"],
            ["Equipment -> ServiceRequests", "1:N", "На одно оборудование может быть создано несколько заявок"],
        ],
        [5.2, 4.0, 8.8],
    )

    heading(doc, "2.3 Проектирование пользовательского интерфейса", 2)
    for text in [
        "Пользовательский интерфейс построен как одно главное окно с вкладками. Такой подход удобен для учебной информационной системы: пользователь не теряется в многооконной навигации и сразу видит основные разделы данных.",
        "Вкладка «Оборудование» содержит таблицу с автоматической генерацией колонок на основе свойств модели Equipment. Вкладка «Клиенты» аналогично отображает данные модели Client. В нижней части окна расположен текстовый блок для вывода сообщений об ошибках подключения.",
        "Окно имеет заголовок «Система обслуживания газового оборудования», ширину 900 и высоту 600 пикселей. Для таблиц используется компонент DataGrid из пакета Avalonia.Controls.DataGrid. Таблицы настроены в режим только для чтения, что соответствует текущему объему функциональности.",
    ]:
        p(doc, text)

    heading(doc, "3 РЕАЛИЗАЦИЯ ПРОГРАММНОГО ПРОДУКТА", 1)
    heading(doc, "3.1 Средства разработки", 2)
    for text in [
        "Разработка выполнена на языке C# с использованием .NET 8. В качестве графического фреймворка выбран Avalonia UI версии 12.0.3. Для доступа к базе данных применяются Microsoft Entity Framework Core 8.0.10 и провайдер Npgsql.EntityFrameworkCore.PostgreSQL 8.0.10.",
        "Выбор указанных технологий обусловлен их распространенностью, поддержкой асинхронной работы и возможностью построения архитектуры MVVM. Entity Framework Core избавляет разработчика от необходимости вручную писать большую часть SQL-запросов для типовых операций выборки и сохранения.",
        "Проект имеет стандартную структуру .NET-приложения: Program.cs отвечает за запуск, App.axaml.cs создает главное окно, папка Views содержит XAML-разметку, ViewModels - модель представления, Models - классы сущностей, Services - сервисы и контекст базы данных.",
    ]:
        p(doc, text)
    add_table(
        doc,
        ["Компонент", "Версия", "Назначение"],
        [
            [".NET", "8.0", "Платформа выполнения и сборки приложения"],
            ["Avalonia", "12.0.3", "Построение настольного пользовательского интерфейса"],
            ["Entity Framework Core", "8.0.10", "ORM для работы с реляционной базой данных"],
            ["Npgsql EF Core Provider", "8.0.10", "Провайдер PostgreSQL для EF Core"],
            ["PostgreSQL", "актуальная локальная версия", "Хранение данных приложения"],
        ],
        [5.0, 3.0, 9.0],
    )

    heading(doc, "3.2 Описание основных модулей", 2)
    for text in [
        "Класс Equipment описывает единицу газового оборудования. Он содержит идентификатор, серийный номер, тип, дату установки, производителя, модель, место установки и текущий статус. Эти свойства напрямую соответствуют колонкам таблицы Equipment.",
        "Класс Client описывает клиента сервисной организации. Для клиента хранятся наименование, контактное лицо, телефон и адрес. Класс ServiceRequest предназначен для дальнейшего развития функциональности заявок и уже отражен в схеме базы данных.",
        "Базовый класс NotifyPropertyChangedBase реализует интерфейс INotifyPropertyChanged. Он нужен для обновления интерфейса при изменении свойств модели представления. Метод SetProperty проверяет, изменилось ли значение, обновляет поле и вызывает событие изменения свойства.",
        "MainViewModel выполняет загрузку данных при создании объекта. Метод LoadDataAsync вызывает сервисные методы GetAllEquipmentAsync и GetAllClientsAsync. При успешной загрузке списки передаются в свойства EquipmentList и ClientList, к которым привязаны таблицы интерфейса.",
        "GasService содержит методы работы с данными. В текущей версии реализованы получение всех записей оборудования и клиентов, а также добавление новых записей в соответствующие наборы DbSet. Сервис отделяет модель представления от деталей работы Entity Framework Core.",
    ]:
        p(doc, text)

    heading(doc, "3.3 Работа с базой данных", 2)
    for text in [
        "Подключение к базе данных настраивается в классе GasDbContext. Строка подключения может быть передана через переменную окружения GAS_SERVICE_DB. Если переменная не задана, используется строка подключения по умолчанию. Такой подход позволяет не изменять исходный код при переносе приложения на другую рабочую станцию.",
        "Метод OnModelCreating задает соответствие классов таблицам базы данных и ограничения длины строковых полей. Для ServiceRequest настроены связи с Equipment и Client через внешние ключи. Удаление связанных записей ограничено правилом Restrict.",
        "SQL-скрипты в папке database позволяют создать базу, схему и начальные данные. В демонстрационном наборе присутствуют реальные для предметной области примеры: газовые котлы, регулятор давления газа, водонагреватель, клиенты из Саратова и Энгельса, а также заявки со статусами «Завершена», «В работе» и «Открыта».",
    ]:
        p(doc, text)

    heading(doc, "4 ТЕСТИРОВАНИЕ И ЭКСПЛУАТАЦИЯ", 1)
    for text in [
        "Проверка программного продукта включала анализ структуры проекта, проверку соответствия моделей схеме базы данных и сборку проекта. Сборка выполнялась командой dotnet build для файла GasServiceApp.csproj.",
        "В результате сборки проект GasServiceApp был успешно скомпилирован в каталог bin/Debug/net8.0. Предупреждения и ошибки компиляции отсутствовали. Это подтверждает корректность ссылок на пакеты, синтаксиса исходного кода и базовой конфигурации проекта.",
        "Функциональное тестирование для текущей версии включает запуск приложения при доступной базе данных, проверку заполнения вкладок «Оборудование» и «Клиенты», а также имитацию ошибки подключения к PostgreSQL. При ошибке пользователь должен увидеть текстовое сообщение в нижней части окна.",
    ]:
        p(doc, text)
    add_table(
        doc,
        ["Тест", "Ожидаемый результат", "Фактический результат"],
        [
            ["Сборка проекта", "Проект компилируется без ошибок", "Выполнено: 0 ошибок, 0 предупреждений"],
            ["Загрузка оборудования", "Таблица отображает записи Equipment", "Реализовано через EquipmentList"],
            ["Загрузка клиентов", "Таблица отображает записи Clients", "Реализовано через ClientList"],
            ["Недоступная БД", "Отображается сообщение об ошибке", "Реализовано через StatusMessage"],
        ],
        [4.5, 6.3, 5.9],
    )
    for text in [
        "Для эксплуатации приложения необходимо установить PostgreSQL, выполнить SQL-скрипты создания базы и начальных данных, затем запустить приложение из среды разработки или из собранного каталога. При отличии параметров подключения от заданных по умолчанию следует настроить переменную окружения GAS_SERVICE_DB.",
        "Дальнейшее развитие системы может включать форму добавления и редактирования оборудования, карточку клиента, модуль регистрации заявок, фильтрацию по статусам, печатные отчеты, авторизацию пользователей и журналирование действий.",
    ]:
        p(doc, text)

    heading(doc, "ЗАКЛЮЧЕНИЕ", 1)
    for text in [
        "В ходе курсовой работы разработано настольное приложение «Система обслуживания газового оборудования». Программа предназначена для учета клиентов и оборудования сервисной организации и использует реляционную базу данных PostgreSQL.",
        "В аналитической части рассмотрена предметная область и сформулированы требования к программному продукту. В проектной части определена архитектура приложения, описана структура базы данных и пользовательский интерфейс. В практической части реализованы модели, сервис доступа к данным, контекст Entity Framework Core и главное окно Avalonia UI.",
        "Поставленная цель достигнута: создано приложение, которое подключается к базе данных, загружает сведения об оборудовании и клиентах и отображает их пользователю в табличном виде. Проект успешно собирается как .NET 8-приложение.",
        "Разработанное решение имеет потенциал развития. На его основе можно реализовать полноценную систему диспетчеризации заявок, контроля сроков обслуживания и формирования управленческой отчетности для газовой сервисной организации.",
    ]:
        p(doc, text)

    heading(doc, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 1)
    sources = [
        "ГОСТ 7.32-2017. Система стандартов по информации, библиотечному и издательскому делу. Отчет о научно-исследовательской работе. Структура и правила оформления. URL: https://protect.gost.ru/default.aspx/document.aspx?control=7&id=218998 (дата обращения: 21.05.2026).",
        "ГОСТ Р 7.0.100-2018. Система стандартов по информации, библиотечному и издательскому делу. Библиографическая запись. Библиографическое описание. URL: https://protect.gost.ru/gost/details/389dfd5a-535a-458a-81c3-14b729b1cee1 (дата обращения: 21.05.2026).",
        "Microsoft. .NET documentation. URL: https://learn.microsoft.com/dotnet/ (дата обращения: 21.05.2026).",
        "Avalonia UI Documentation. URL: https://docs.avaloniaui.net/ (дата обращения: 21.05.2026).",
        "Microsoft. Entity Framework Core documentation. URL: https://learn.microsoft.com/ef/core/ (дата обращения: 21.05.2026).",
        "PostgreSQL Documentation. URL: https://www.postgresql.org/docs/ (дата обращения: 21.05.2026).",
        "Npgsql Entity Framework Core Provider documentation. URL: https://www.npgsql.org/efcore/ (дата обращения: 21.05.2026).",
    ]
    for i, source in enumerate(sources, 1):
        p(doc, f"{i}. {source}", align=WD_ALIGN_PARAGRAPH.LEFT)

    doc.add_page_break()
    heading(doc, "ПРИЛОЖЕНИЕ А", 1)
    caption(doc, "Фрагменты программного кода")
    add_code_block(
        doc,
        "Листинг А.1 - Модель оборудования",
        """
public class Equipment
{
    public int Id { get; set; }
    public string SerialNumber { get; set; } = string.Empty;
    public string Type { get; set; } = string.Empty;
    public DateTime InstallationDate { get; set; }
    public string Manufacturer { get; set; } = string.Empty;
    public string Model { get; set; } = string.Empty;
    public string Location { get; set; } = string.Empty;
    public string Status { get; set; } = string.Empty;
}
""",
    )
    add_code_block(
        doc,
        "Листинг А.2 - Загрузка данных в MainViewModel",
        """
private async Task LoadDataAsync()
{
    try
    {
        EquipmentList = await _service.GetAllEquipmentAsync();
        ClientList = await _service.GetAllClientsAsync();
        StatusMessage = string.Empty;
    }
    catch (Exception ex)
    {
        StatusMessage = $"Ошибка подключения к PostgreSQL: {ex.Message}";
    }
}
""",
    )
    add_code_block(
        doc,
        "Листинг А.3 - Методы сервиса",
        """
public async Task<List<Equipment>> GetAllEquipmentAsync()
{
    return await _context.Equipment.ToListAsync();
}

public async Task<List<Client>> GetAllClientsAsync()
{
    return await _context.Clients.ToListAsync();
}
""",
    )

    doc.add_page_break()
    heading(doc, "ПРИЛОЖЕНИЕ Б", 1)
    caption(doc, "SQL-скрипт структуры базы данных")
    add_code_block(
        doc,
        "Листинг Б.1 - Основные таблицы базы данных",
        """
CREATE TABLE IF NOT EXISTS "Equipment" (
    "Id" integer GENERATED BY DEFAULT AS IDENTITY PRIMARY KEY,
    "SerialNumber" varchar(64) NOT NULL,
    "Type" varchar(100) NOT NULL,
    "InstallationDate" timestamp without time zone NOT NULL,
    "Manufacturer" varchar(100) NOT NULL,
    "Model" varchar(100) NOT NULL,
    "Location" varchar(200) NOT NULL,
    "Status" varchar(50) NOT NULL
);

CREATE TABLE IF NOT EXISTS "Clients" (
    "Id" integer GENERATED BY DEFAULT AS IDENTITY PRIMARY KEY,
    "Name" varchar(200) NOT NULL,
    "ContactPerson" varchar(150) NOT NULL,
    "Phone" varchar(30) NOT NULL,
    "Address" varchar(250) NOT NULL
);

CREATE TABLE IF NOT EXISTS "ServiceRequests" (
    "Id" integer GENERATED BY DEFAULT AS IDENTITY PRIMARY KEY,
    "EquipmentId" integer NOT NULL,
    "ClientId" integer NOT NULL,
    "RequestDate" timestamp without time zone NOT NULL,
    "CompletionDate" timestamp without time zone NULL,
    "Status" varchar(50) NOT NULL,
    "Description" varchar(1000) NOT NULL
);
""",
    )


def main():
    doc = setup_document()
    add_title_page(doc)
    add_contents(doc)
    build_body(doc)
    doc.save(OUT)


if __name__ == "__main__":
    main()
